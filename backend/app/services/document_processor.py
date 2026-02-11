"""Document text extraction for various file types."""

import csv
import io
from pathlib import Path

from app.models.knowledge import FileType


def extract_text(file_path: str, file_type: FileType) -> str:
    """Extract text content from a document file.

    Args:
        file_path: Path to the file on disk
        file_type: Type of file (pdf, docx, txt, csv, json)

    Returns:
        Extracted text content as a string
    """
    path = Path(file_path)

    if file_type == FileType.txt:
        return path.read_text(encoding="utf-8")

    if file_type == FileType.csv:
        return _extract_csv(path)

    if file_type == FileType.json:
        # Read JSON as text (for knowledge base purposes, treat as plain text)
        return path.read_text(encoding="utf-8")

    if file_type == FileType.pdf:
        return _extract_pdf(path)

    if file_type == FileType.docx:
        return _extract_docx(path)

    raise ValueError(f"Unsupported file type: {file_type}")


def _extract_csv(path: Path) -> str:
    """Convert CSV to readable text format."""
    text_parts = []
    with path.open(encoding="utf-8") as f:
        reader = csv.reader(f)
        headers = next(reader, None)
        if headers:
            text_parts.append("Headers: " + ", ".join(headers))
            for row in reader:
                row_text = ", ".join(
                    f"{h}: {v}" for h, v in zip(headers, row) if v.strip()
                )
                if row_text:
                    text_parts.append(row_text)
    return "\n".join(text_parts)


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF extraction. Install with: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx (paragraphs and tables)."""
    if not path.exists():
        raise FileNotFoundError(f"Document file not found: {path}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. Install with: pip install python-docx"
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValueError(f"Cannot open DOCX (corrupted or invalid format): {e}") from e

    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    result = "\n\n".join(text_parts)
    if not result.strip():
        raise ValueError(
            "No text could be extracted from the DOCX (document may contain only images or be empty)."
        )
    return result
